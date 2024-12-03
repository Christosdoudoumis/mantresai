import logging
from flask import Flask, request, jsonify, render_template_string, redirect, url_for, send_file
import requests
from bs4 import BeautifulSoup
from groq import Groq
import validators
from io import BytesIO
from docx import Document
from datetime import datetime
import uuid
import webbrowser
import threading

app = Flask(__name__)
app.secret_key = 'your_secure_secret_key'  # Replace with a secure key

# Configure Logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)

# Hardcoded Groq API Key
GROQ_API_KEY = "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"  # Replace with your actual key

# Groq client
try:
    client = Groq(api_key=GROQ_API_KEY)
except Exception as e:
    logging.error(f"Failed to initialize Groq client: {e}")
    raise

# Constants
MAX_CHARACTERS = 5000  # Adjust based on model limits
REQUEST_TIMEOUT = 10   # Timeout for HTTP requests in seconds

# In-memory storage for conversations
# In a production environment, consider using a database
conversations = {}

# HTML Templates
CHATBOT_UI = '''
<!doctype html>
<html lang="el">
<head>
    <meta charset="utf-8">
    <title>Βοηθός Εξαγωγής Δεδομένων</title>
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <!-- Bootstrap CSS -->
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <!-- Font Awesome for Icons -->
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.5.0/css/all.min.css" integrity="sha512-SomeHashValue" crossorigin="anonymous" referrerpolicy="no-referrer" />
    <!-- Google Fonts -->
    <link href="https://fonts.googleapis.com/css2?family=Nunito:wght@300;400;600;700&display=swap" rel="stylesheet">
    <style>
        /* Global Styles */
        body {
            background: #f0f2f5;
            font-family: 'Nunito', sans-serif;
            height: 100vh;
            margin: 0;
            display: flex;
            align-items: center;
            justify-content: center;
            overflow: hidden;
        }

        /* Chat Container */
        .chat-container {
            background-color: #ffffff;
            width: 100%;
            max-width: 800px;
            height: 90vh;
            display: flex;
            flex-direction: column;
            border-radius: 15px;
            box-shadow: 0 10px 25px rgba(0, 0, 0, 0.1);
            overflow: hidden;
        }

        /* Chat Header */
        .chat-header {
            background-color: #0d6efd;
            color: #fff;
            padding: 20px;
            display: flex;
            align-items: center;
            justify-content: space-between;
        }

        .chat-header .logo {
            display: flex;
            align-items: center;
        }

        .chat-header .logo i {
            font-size: 1.8rem;
            margin-right: 10px;
        }

        .chat-header h1 {
            font-size: 1.5rem;
            margin: 0;
            font-weight: 700;
        }

        /* Navigation Links */
        .nav-links {
            display: flex;
            gap: 20px;
        }

        .nav-links a {
            color: #fff;
            text-decoration: none;
            font-weight: 500;
            position: relative;
            padding-bottom: 5px;
        }

        .nav-links a::after {
            content: '';
            position: absolute;
            width: 0%;
            height: 2px;
            bottom: 0;
            left: 0;
            background-color: #fff;
            transition: width 0.3s;
        }

        .nav-links a:hover::after {
            width: 100%;
        }

        /* Chat Messages */
        .chat-messages {
            flex: 1;
            padding: 20px;
            overflow-y: auto;
            background: #f8f9fa;
        }

        .chat-message {
            display: flex;
            margin-bottom: 20px;
            align-items: flex-end;
            animation: fadeIn 0.5s ease-in-out;
        }

        .chat-message.user {
            justify-content: flex-end;
        }

        .chat-message.bot {
            justify-content: flex-start;
        }

        .message-content {
            max-width: 70%;
            padding: 15px 20px;
            border-radius: 20px;
            position: relative;
            font-size: 1rem;
            line-height: 1.6;
            background-color: #e9ecef;
            color: #333;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }

        .chat-message.user .message-content {
            background-color: #0d6efd;
            color: #fff;
            border-bottom-right-radius: 0;
        }

        .chat-message.bot .message-content {
            background-color: #e9ecef;
            color: #333;
            border-bottom-left-radius: 0;
        }

        .chat-message .avatar {
            width: 40px;
            height: 40px;
            border-radius: 50%;
            background-color: #6c757d;
            display: flex;
            justify-content: center;
            align-items: center;
            color: #fff;
            font-size: 1.2rem;
            margin: 0 10px;
            flex-shrink: 0;
        }

        /* Chat Footer */
        .chat-footer {
            padding: 20px;
            background-color: #ffffff;
            border-top: 1px solid #dee2e6;
        }

        .chat-footer form {
            display: flex;
            align-items: center;
        }

        .chat-footer input[type="text"] {
            flex: 1;
            padding: 12px 20px;
            border: 1px solid #ced4da;
            border-radius: 25px;
            font-size: 1rem;
            outline: none;
            transition: border-color 0.2s, box-shadow 0.2s;
        }

        .chat-footer input[type="text"]:focus {
            border-color: #0d6efd;
            box-shadow: 0 0 8px rgba(13, 110, 253, 0.3);
        }

        .chat-footer button {
            margin-left: 15px;
            padding: 12px 16px;
            background-color: #0d6efd;
            color: #fff;
            border: none;
            border-radius: 50%;
            cursor: pointer;
            font-size: 1.2rem;
            transition: background-color 0.3s, transform 0.2s;
            display: flex;
            align-items: center;
            justify-content: center;
        }

        .chat-footer button:hover {
            background-color: #0b5ed7;
            transform: scale(1.05);
        }

        /* Scrollbar Styling */
        .chat-messages::-webkit-scrollbar {
            width: 8px;
        }

        .chat-messages::-webkit-scrollbar-thumb {
            background: #ced4da;
            border-radius: 4px;
        }

        .chat-messages::-webkit-scrollbar-thumb:hover {
            background: #adb5bd;
        }

        /* Animations */
        @keyframes fadeIn {
            from { opacity: 0; transform: translateY(10px); }
            to { opacity: 1; transform: translateY(0); }
        }

        /* Typing Indicator */
        .typing-indicator {
            display: flex;
            align-items: center;
            gap: 5px;
        }

        .typing-indicator .dot {
            width: 8px;
            height: 8px;
            background-color: #6c757d;
            border-radius: 50%;
            animation: bounce 1.4s infinite;
        }

        .typing-indicator .dot:nth-child(1) {
            animation-delay: -0.32s;
        }

        .typing-indicator .dot:nth-child(2) {
            animation-delay: -0.16s;
        }

        .typing-indicator .dot:nth-child(3) {
            animation-delay: 0s;
        }

        @keyframes bounce {
            0%, 80%, 100% { transform: scale(0); }
            40% { transform: scale(1); }
        }

        /* Responsive Styles */
        @media (max-width: 768px) {
            .chat-container {
                height: 100vh;
                border-radius: 0;
            }
            .chat-header h1 {
                font-size: 1.3rem;
            }
            .chat-footer input[type="text"] {
                font-size: 0.9rem;
                padding: 10px 15px;
            }
            .chat-footer button {
                font-size: 1rem;
                padding: 10px 14px;
                margin-left: 10px;
            }
            .message-content {
                font-size: 0.95rem;
                padding: 12px 18px;
            }
            .chat-message .avatar {
                width: 35px;
                height: 35px;
                font-size: 1rem;
            }
        }

        @media (max-width: 576px) {
            .nav-links {
                display: none; /* Hide navigation links on very small screens */
            }
            .chat-header {
                padding: 15px;
            }
            .chat-footer input[type="text"] {
                padding: 8px 12px;
            }
            .chat-footer button {
                margin-left: 8px;
            }
        }

        /* Dark Mode Toggle */
        .dark-mode-toggle {
            cursor: pointer;
            font-size: 1.2rem;
            transition: color 0.3s;
        }

        body.dark-mode {
            background: #121212;
            color: #ffffff;
        }

        body.dark-mode .chat-container {
            background-color: #1e1e1e;
            box-shadow: 0 10px 25px rgba(255, 255, 255, 0.1);
        }

        body.dark-mode .chat-header {
            background-color: #333333;
        }

        body.dark-mode .nav-links a {
            color: #ffffff;
        }

        body.dark-mode .message-content {
            background-color: #333333;
            color: #ffffff;
        }

        body.dark-mode .chat-message.user .message-content {
            background-color: #007bff;
        }

        body.dark-mode .chat-footer {
            background-color: #1e1e1e;
            border-top: 1px solid #444444;
        }

        body.dark-mode .chat-footer input[type="text"] {
            background-color: #2c2c2c;
            border: 1px solid #555555;
            color: #ffffff;
        }

        body.dark-mode .chat-footer input[type="text"]::placeholder {
            color: #aaaaaa;
        }

        body.dark-mode .chat-footer button {
            background-color: #007bff;
        }

        body.dark-mode .chat-footer button:hover {
            background-color: #0056b3;
        }

        body.dark-mode .avatar {
            background-color: #555555;
        }

        body.dark-mode .typing-indicator .dot {
            background-color: #cccccc;
        }

    </style>
</head>
<body>
    <div class="chat-container">
        <!-- Chat Header -->
        <div class="chat-header">
            <div class="logo">
                <i class="fas fa-car"></i>
                <h1>Βοηθός Εξαγωγής Δεδομένων</h1>
            </div>
            <div class="nav-links">
                <a href="{{ url_for('index') }}">Συνομιλία</a>
                <a href="{{ url_for('history') }}">Ιστορικό</a>
                <a href="javascript:void(0);" id="dark-mode-toggle" title="Εναλλαγή Σκούρου/Φωτεινού Θέματος">
                    <i class="fas fa-moon dark-mode-toggle"></i>
                </a>
            </div>
        </div>

        <!-- Chat Messages -->
        <div class="chat-messages" id="chat-messages">
            <!-- Messages will be appended here -->
            {% if conversation %}
                {% for msg in conversation %}
                    <div class="chat-message {{ msg.sender }}">
                        {% if msg.sender == 'bot' %}
                            <div class="avatar"><i class="fas fa-robot"></i></div>
                            <div class="message-content">{{ msg.message|safe }}</div>
                        {% else %}
                            <div class="message-content">{{ msg.message|safe }}</div>
                            <div class="avatar"><i class="fas fa-user"></i></div>
                        {% endif %}
                    </div>
                {% endfor %}
            {% endif %}
        </div>

        <!-- Chat Footer -->
        {% if not conversation %}
        <div class="chat-footer">
            <form id="chat-form" onsubmit="return sendMessage(event);">
                <input type="text" id="user-input" placeholder="Εισάγετε μια URL για ανάλυση..." required autocomplete="off">
                <button type="submit"><i class="fas fa-paper-plane"></i></button>
            </form>
        </div>
        {% endif %}
    </div>

    <!-- Bootstrap JS Bundle with Popper -->
    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
    <!-- Font Awesome JS (for icons in dynamically added content) -->
    <script src="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.5.0/js/all.min.js" integrity="sha512-SomeHashValue" crossorigin="anonymous" referrerpolicy="no-referrer"></script>
    {% if not conversation %}
    <script>
        const chatMessages = document.getElementById('chat-messages');
        const darkModeToggle = document.getElementById('dark-mode-toggle');

        // Dark Mode Toggle Functionality
        darkModeToggle.addEventListener('click', () => {
            document.body.classList.toggle('dark-mode');
            const icon = darkModeToggle.querySelector('i');
            icon.classList.toggle('fa-moon');
            icon.classList.toggle('fa-sun');
        });

        function appendMessage(sender, message) {
            const messageElement = document.createElement('div');
            messageElement.classList.add('chat-message', sender);

            const avatar = document.createElement('div');
            avatar.classList.add('avatar');
            avatar.innerHTML = sender === 'user' ? '<i class="fas fa-user"></i>' : '<i class="fas fa-robot"></i>';

            const messageContent = document.createElement('div');
            messageContent.classList.add('message-content');

            if (message === 'typing-indicator') {
                messageContent.innerHTML = `
                    <div class="typing-indicator">
                        <div class="dot"></div>
                        <div class="dot"></div>
                        <div class="dot"></div>
                    </div>
                `;
            } else {
                messageContent.innerHTML = message;
            }

            if (sender === 'user') {
                messageElement.appendChild(messageContent);
                messageElement.appendChild(avatar);
            } else {
                messageElement.appendChild(avatar);
                messageElement.appendChild(messageContent);
            }

            chatMessages.appendChild(messageElement);
            chatMessages.scrollTop = chatMessages.scrollHeight;
        }

        async function sendMessage(event) {
            event.preventDefault();
            const userInput = document.getElementById('user-input');
            const message = userInput.value.trim();
            if (!message) return;
            appendMessage('user', message);
            userInput.value = '';

            // Append typing indicator
            appendMessage('bot', 'typing-indicator');

            try {
                const response = await fetch('/chat', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ message })
                });

                if (!response.ok) {
                    const errorData = await response.json();
                    throw new Error(errorData.response || "Λυπάμαι, κάτι πήγε στραβά.");
                }

                const data = await response.json();

                // Remove typing indicator
                const lastMessage = chatMessages.lastChild;
                if (lastMessage && lastMessage.querySelector('.typing-indicator')) {
                    chatMessages.removeChild(lastMessage);
                }

                appendMessage('bot', data.response || "Λυπάμαι, κάτι πήγε στραβά.");
            } catch (error) {
                // Remove typing indicator
                const lastMessage = chatMessages.lastChild;
                if (lastMessage && lastMessage.querySelector('.typing-indicator')) {
                    chatMessages.removeChild(lastMessage);
                }
                appendMessage('bot', error.message || "Λυπάμαι, κάτι πήγε στραβά.");
            }
        }
    </script>
    {% endif %}
</body>
</html>
'''

HISTORY_UI = '''
<!doctype html>
<html lang="el">
<head>
    <meta charset="utf-8">
    <title>Ιστορικό Συνομιλιών</title>
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <!-- Bootstrap CSS -->
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <!-- Font Awesome for Icons -->
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.5.0/css/all.min.css" integrity="sha512-SomeHashValue" crossorigin="anonymous" referrerpolicy="no-referrer" />
    <!-- Google Fonts -->
    <link href="https://fonts.googleapis.com/css2?family=Nunito:wght@300;400;600;700&display=swap" rel="stylesheet">
    <style>
        body {
            background: #f0f2f5;
            font-family: 'Nunito', sans-serif;
            padding: 20px;
        }

        .history-container {
            background-color: #ffffff;
            padding: 30px;
            border-radius: 15px;
            box-shadow: 0 10px 25px rgba(0, 0, 0, 0.1);
            max-width: 1000px;
            margin: auto;
        }

        .history-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 30px;
        }

        .history-header h2 {
            margin: 0;
            font-weight: 700;
            color: #0d6efd;
        }

        .conversation {
            border: 1px solid #dee2e6;
            border-radius: 10px;
            padding: 20px;
            margin-bottom: 25px;
            background-color: #f8f9fa;
            transition: background-color 0.3s;
        }

        .conversation:hover {
            background-color: #e2e6ea;
        }

        .conversation-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 15px;
        }

        .conversation-header span {
            font-size: 0.95rem;
            color: #6c757d;
        }

        .conversation-messages {
            max-height: 300px;
            overflow-y: auto;
            margin-bottom: 15px;
        }

        .chat-message {
            display: flex;
            margin-bottom: 15px;
            align-items: flex-end;
        }

        .chat-message.user {
            justify-content: flex-end;
        }

        .chat-message.bot {
            justify-content: flex-start;
        }

        .message-content {
            max-width: 70%;
            padding: 10px 15px;
            border-radius: 15px;
            font-size: 0.9rem;
            line-height: 1.4;
            background-color: #e9ecef;
            color: #333;
            box-shadow: 0 2px 5px rgba(0,0,0,0.05);
        }

        .chat-message.user .message-content {
            background-color: #0d6efd;
            color: #fff;
            border-bottom-right-radius: 0;
        }

        .chat-message.bot .message-content {
            background-color: #e9ecef;
            color: #333;
            border-bottom-left-radius: 0;
        }

        .avatar {
            width: 30px;
            height: 30px;
            border-radius: 50%;
            background-color: #6c757d;
            display: flex;
            justify-content: center;
            align-items: center;
            color: #fff;
            font-size: 0.9rem;
            margin: 0 8px;
            flex-shrink: 0;
        }

        .download-button {
            background-color: #0d6efd;
            color: #fff;
            border: none;
            padding: 8px 12px;
            border-radius: 5px;
            cursor: pointer;
            font-size: 0.9rem;
            transition: background-color 0.3s;
            text-decoration: none;
            display: flex;
            align-items: center;
            gap: 5px;
        }

        .download-button:hover {
            background-color: #0b5ed7;
        }

        /* Responsive Styles */
        @media (max-width: 768px) {
            .history-container {
                padding: 20px;
            }
            .conversation-messages {
                max-height: 250px;
            }
        }

        @media (max-width: 576px) {
            .history-header h2 {
                font-size: 1.2rem;
            }
            .download-button {
                padding: 6px 10px;
                font-size: 0.8rem;
            }
        }
    </style>
</head>
<body>
    <div class="history-container">
        <div class="history-header">
            <h2>Ιστορικό Συνομιλιών</h2>
            <a href="{{ url_for('index') }}" class="btn btn-primary"><i class="fas fa-arrow-left"></i> Επιστροφή στη Συνομιλία</a>
        </div>
        {% if all_conversations %}
            {% for conv_id, conv in all_conversations.items() %}
                <div class="conversation">
                    <div class="conversation-header">
                        <span><strong>Συνομιλία:</strong> {{ conv['timestamp'] }}</span>
                        <a href="{{ url_for('download_convo', convo_id=conv_id) }}" class="download-button"><i class="fas fa-download"></i> Κατέβασμα DOCX</a>
                    </div>
                    <div class="conversation-messages">
                        {% for msg in conv['messages'] %}
                            <div class="chat-message {{ msg.sender }}">
                                {% if msg.sender == 'bot' %}
                                    <div class="avatar"><i class="fas fa-robot"></i></div>
                                    <div class="message-content">{{ msg.message|safe }}</div>
                                {% else %}
                                    <div class="message-content">{{ msg.message|safe }}</div>
                                    <div class="avatar"><i class="fas fa-user"></i></div>
                                {% endif %}
                            </div>
                        {% endfor %}
                    </div>
                </div>
            {% endfor %}
        {% else %}
            <div class="text-center">
                <p>Δεν υπάρχουν συνομιλίες στο ιστορικό.</p>
            </div>
        {% endif %}
    </div>

    <!-- Bootstrap JS Bundle with Popper -->
    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
    <!-- Font Awesome JS -->
    <script src="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.5.0/js/all.min.js" integrity="sha512-SomeHashValue" crossorigin="anonymous" referrerpolicy="no-referrer"></script>
</body>
</html>
'''

@app.route('/')
def index():
    return render_template_string(CHATBOT_UI, conversation=None)

@app.route('/history')
def history():
    return render_template_string(HISTORY_UI, all_conversations=conversations)

@app.route('/chat', methods=['POST'])
def chat():
    user_message = request.json.get('message', '').strip()
    if not user_message:
        logging.warning("Empty message received.")
        return jsonify({'response': "Δεν έλαβα κάποιο μήνυμα. Παρακαλώ εισάγετε μια URL για ανάλυση."}), 400

    logging.info(f"Received user message: {user_message}")

    # Validate URL
    if not validators.url(user_message):
        logging.warning(f"Invalid URL provided: {user_message}")
        return jsonify({'response': "Παρακαλώ εισάγετε μια έγκυρη URL."}), 400

    try:
        logging.info(f"Fetching content from URL: {user_message}")
        response = requests.get(user_message, timeout=REQUEST_TIMEOUT)
        response.raise_for_status()
        response.encoding = 'utf-8'
        webpage_content = response.text

        # Parse and extract text
        soup = BeautifulSoup(webpage_content, 'html.parser')
        text_content = soup.get_text(separator='\n', strip=True)

        # Limit content size
        if len(text_content) > MAX_CHARACTERS:
            logging.info("Content is too large; truncating.")
            text_content = text_content[:MAX_CHARACTERS] + "..."

        # Prepare prompt for LLM
        prompt = f"""
Ο χρήστης θέλει να λάβει τις εξής πληροφορίες από την ιστοσελίδα {user_message}. Παρέδωσε τα δεδομένα με ευανάγνωστο τρόπο, ομαδοποιημένα ανά κατηγορία, και παρουσίασέ τα ως εξής:

---
**1. Διευθύνσεις Email**  
Παρουσίασε όλες τις email διευθύνσεις που εντοπίστηκαν στη σελίδα, μία ανά γραμμή.

---
**2. Τηλεφωνικοί Αριθμοί**  
Κατέγραψε όλους τους τηλεφωνικούς αριθμούς που εμφανίζονται στη σελίδα, ομαδοποιώντας τους αν υπάρχουν διαφορετικοί τύποι (π.χ., κινητά, σταθερά, διεθνείς).

---
**3. Διευθύνσεις**  
Παρέδωσε τις φυσικές διευθύνσεις που εντοπίστηκαν στη σελίδα, ταξινομημένες αν είναι δυνατόν (π.χ., οδός, αριθμός, πόλη, ταχυδρομικός κώδικας).

---
**4. Περιγραφή Σελίδας**  
Παρουσίασε μια σύντομη περιγραφή του περιεχομένου της σελίδας στα Ελληνικά:  
- Αν υπάρχει meta περιγραφή, χρησιμοποίησέ την.  
- Αν όχι, δημιούργησε μια σύντομη περίληψη του κειμένου που εμφανίζεται στη σελίδα.

---
Φρόντισε η παρουσίαση να είναι καθαρή, πλήρης και χωρίς περιττές πληροφορίες.

Εδώ είναι το περιεχόμενο της ιστοσελίδας:

{text_content}
"""

        logging.info("Sending prompt to LLM.")
        chat_completion = client.chat.completions.create(
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": ""}
            ],
            model="llama3-8b-8192"
        )

        response_text = chat_completion.choices[0].message.content.strip()
        logging.info("Received response from LLM.")

        # Store conversation
        convo_id = str(uuid.uuid4())
        conversations[convo_id] = {
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'messages': [
                {'sender': 'user', 'message': user_message},
                {'sender': 'bot', 'message': response_text}
            ]
        }

        return jsonify({'response': response_text})

    except requests.exceptions.Timeout:
        logging.error("Request timed out.")
        return jsonify({'response': "Η αίτηση στον εξωτερικό διακομιστή ξόδεψε. Παρακαλώ δοκιμάστε ξανά αργότερα."}), 504
    except requests.exceptions.HTTPError as http_err:
        logging.error(f"HTTP error occurred: {http_err}")
        return jsonify({'response': f"Σφάλμα HTTP κατά την πρόσβαση στην ιστοσελίδα: {str(http_err)}"}), 502
    except requests.exceptions.RequestException as req_err:
        logging.error(f"Request exception: {req_err}")
        return jsonify({'response': f"Σφάλμα κατά την προσπάθεια πρόσβασης στην ιστοσελίδα: {str(req_err)}"}), 502
    except AttributeError as attr_err:
        logging.error(f"Attribute error: {attr_err}")
        return jsonify({'response': f"Σφάλμα στην επεξεργασία της απάντησης από το LLM: {str(attr_err)}"}), 500
    except Exception as e:
        logging.error(f"Unexpected error: {e}")
        return jsonify({'response': f"Απρόσμενο σφάλμα: {str(e)}"}), 500

@app.route('/download/<convo_id>')
def download_convo(convo_id):
    convo = conversations.get(convo_id)
    if not convo:
        logging.warning(f"Conversation ID {convo_id} not found.")
        return "Conversation not found.", 404

    try:
        # Create a DOCX document
        doc = Document()
        doc.add_heading(f'Συνομιλία από {convo["timestamp"]}', 0)

        for msg in convo['messages']:
            if msg['sender'] == 'user':
                doc.add_paragraph('Χρήστης:', style='Intense Quote')
                doc.add_paragraph(msg['message'])
            else:
                doc.add_paragraph('Βοηθός:', style='Intense Quote')
                doc.add_paragraph(msg['message'])

        # Save the document to a BytesIO stream
        byte_io = BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)

        filename = f'Conversation_{convo_id}.docx'

        return send_file(
            byte_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logging.error(f"Error generating DOCX: {e}")
        return "Σφάλμα κατά τη δημιουργία του αρχείου DOCX.", 500

def open_browser():
    webbrowser.open_new("http://localhost:5000/")

if __name__ == '__main__':
    # Automatically open the default web browser after 1.25 seconds
    threading.Timer(1.25, open_browser).start()
    # Run the Flask app
    app.run(debug=False, host='0.0.0.0', port=5000)
